import json
from datetime import datetime
from typing import Optional
from fastapi import APIRouter, Depends, HTTPException, Request
from pydantic import BaseModel
from sqlalchemy.orm import Session
from database import get_db
from models import FlowRequest
import audit
from engine.validator import validate_flow
from engine.path_finder import find_path
from engine.script_generator import generate_scripts
from compliance import default_provider, build_flow_subject

router = APIRouter()


class StatusUpdate(BaseModel):
    status: str  # pending, validated, deployed, rejected
    rejection_reason: Optional[str] = None


class FlowIn(BaseModel):
    src_ip: str
    dst_ip: str
    port: str
    protocol: str = "tcp"
    application: str = ""
    justification: str = ""
    analyst: str = "demo-user"


@router.post("/analyze")
def analyze_flow(data: FlowIn, compliance_source: Optional[str] = None,
                 db: Session = Depends(get_db)):
    """Analyse un flux sans le persister — pour l'aperçu temps réel.

    Inclut le verdict de conformité (moteur OSCAL) évalué sur le chemin réel.
    `compliance_source` permet de cibler une source (nis2/anssi-hygiene/cis-v8) ;
    par défaut tout le catalogue est évalué."""
    validation = validate_flow(data.src_ip, data.dst_ip, data.port, data.protocol, db)
    path = {}
    scripts = {}
    compliance = None

    # Chemin calculé indépendamment du verdict (topologique) — sert au rendu
    # ET à la conformité, évaluée même sur un flux rejeté tant que les zones
    # sont résolues.
    comp_path = find_path(data.src_ip, data.dst_ip, db)

    if validation["valid"]:
        path = comp_path
        if path.get("found"):
            scripts = generate_scripts(
                src_ip=data.src_ip,
                dst_ip=data.dst_ip,
                port=data.port,
                protocol=data.protocol,
                application=data.application,
                justification=data.justification,
                src_zone=validation.get("src_zone"),
                dst_zone=validation.get("dst_zone"),
                path_hops=path.get("hops", []),
            )

    if validation.get("src_zone") and validation.get("dst_zone"):
        subject = build_flow_subject(
            db, data.src_ip, data.dst_ip, data.port, data.protocol,
            validation.get("src_zone"), validation.get("dst_zone"),
            comp_path.get("hops", []),
        )
        try:
            compliance = default_provider.evaluate(subject, source=compliance_source).to_dict()
        except KeyError:
            compliance = None

    return {"validation": validation, "path": path, "scripts": scripts, "compliance": compliance}


@router.post("", status_code=201)
def submit_flow(data: FlowIn, request: Request = None, db: Session = Depends(get_db)):
    """Soumet et persiste une demande de flux."""
    validation = validate_flow(data.src_ip, data.dst_ip, data.port, data.protocol, db)
    path = {}
    scripts = {}

    if validation["valid"]:
        path = find_path(data.src_ip, data.dst_ip, db)
        if path.get("found"):
            scripts = generate_scripts(
                src_ip=data.src_ip,
                dst_ip=data.dst_ip,
                port=data.port,
                protocol=data.protocol,
                application=data.application,
                justification=data.justification,
                src_zone=validation.get("src_zone"),
                dst_zone=validation.get("dst_zone"),
                path_hops=path.get("hops", []),
            )

    status = "pending"
    flow = FlowRequest(
        src_ip=data.src_ip,
        dst_ip=data.dst_ip,
        port=data.port,
        protocol=data.protocol,
        application=data.application,
        justification=data.justification,
        analyst=data.analyst,
        status=status,
        validation_result=json.dumps(validation),
        path_result=json.dumps(path),
        scripts_result=json.dumps(scripts),
    )
    db.add(flow)
    db.commit()
    db.refresh(flow)

    audit.record_audit(
        db, action="CREATE", object_type="FLOW",
        object_id=f"FLOW-{flow.id}",
        object_name=f"{data.src_ip} → {data.dst_ip}:{data.port}",
        application=data.application or None,
        details={"port": data.port, "protocol": data.protocol, "valid": validation.get("valid")},
        after={"status": status, "src_ip": data.src_ip, "dst_ip": data.dst_ip, "port": data.port},
        user_id=data.analyst or "admin", username=data.analyst or "admin",
        ip_address=audit.client_ip(request),
    )

    return {
        "id": flow.id,
        "status": status,
        "validation": validation,
        "path": path,
        "scripts": scripts,
    }


@router.get("")
def list_flows(db: Session = Depends(get_db)):
    flows = db.query(FlowRequest).order_by(FlowRequest.created_at.desc()).all()
    result = []
    for f in flows:
        result.append({
            "id": f.id,
            "created_at": f.created_at.isoformat(),
            "src_ip": f.src_ip,
            "dst_ip": f.dst_ip,
            "port": f.port,
            "protocol": f.protocol,
            "application": f.application,
            "status": f.status,
            "analyst": f.analyst,
        })
    return result


@router.get("/export/docx")
def export_flows_docx(
    columns: str = "id,date,src_ip,dst_ip,port,protocol,application,analyst,status",
    statuses: Optional[str] = None,
    search: Optional[str] = None,
    db: Session = Depends(get_db),
):
    """Génère un fichier Word (.docx) avec la liste des flux filtrés."""
    import io
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from fastapi.responses import Response

    STATUS_LABELS = {
        "validated": "Validé", "deployed": "Déployé",
        "rejected": "Refusé", "pending": "En attente",
    }
    COLUMN_DEFS = {
        "id":            ("ID",             lambda f: str(f.id)),
        "date":          ("Date",           lambda f: f.created_at.strftime("%d/%m/%Y")),
        "src_ip":        ("IP Source",      lambda f: f.src_ip or ""),
        "dst_ip":        ("IP Destination", lambda f: f.dst_ip or ""),
        "port":          ("Port",           lambda f: str(f.port or "")),
        "protocol":      ("Protocole",      lambda f: (f.protocol or "").upper()),
        "application":   ("Application",    lambda f: f.application or "—"),
        "analyst":       ("Analyste",       lambda f: f.analyst or "—"),
        "status":        ("Statut",         lambda f: STATUS_LABELS.get(f.status, f.status)),
        "justification": ("Justification",  lambda f: f.justification or "—"),
    }

    flows_q = db.query(FlowRequest).order_by(FlowRequest.created_at.desc()).all()
    status_list = [s.strip() for s in statuses.split(",")] if statuses else []
    if status_list:
        flows_q = [f for f in flows_q if f.status in status_list]
    if search:
        s = search.lower()
        flows_q = [
            f for f in flows_q
            if s in (f.src_ip or "") or s in (f.dst_ip or "")
            or s in (f.application or "").lower()
            or s in (f.analyst or "").lower()
        ]

    col_keys = [c.strip() for c in columns.split(",")]
    active_cols = [(k, COLUMN_DEFS[k]) for k in col_keys if k in COLUMN_DEFS]

    doc = Document()
    for section in doc.sections:
        section.page_width  = Cm(29.7)
        section.page_height = Cm(21.0)
        section.left_margin   = Cm(1.5)
        section.right_margin  = Cm(1.5)
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)

    # ── Page de garde ──────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Rapport — Matrice des flux IP")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run(f"Généré le {datetime.now().strftime('%d/%m/%Y à %Hh%M')}")

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run(f"{len(flows_q)} flux exportés")

    if status_list:
        p4 = doc.add_paragraph()
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        labels = ", ".join(STATUS_LABELS.get(s, s) for s in status_list)
        p4.add_run(f"Filtre statut : {labels}")
    if search:
        p5 = doc.add_paragraph()
        p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p5.add_run(f'Recherche : "{search}"')

    doc.add_page_break()

    # ── Tableau ─────────────────────────────────────────────────────────────
    def set_cell_bg(cell, hex_color: str):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tcPr.append(shd)

    if not active_cols:
        doc.add_paragraph("Aucune colonne sélectionnée.")
    else:
        table = doc.add_table(rows=1, cols=len(active_cols))
        table.style = "Table Grid"

        # En-têtes
        hdr = table.rows[0].cells
        for i, (_, (label, _)) in enumerate(active_cols):
            cell = hdr[i]
            cell.text = ""
            set_cell_bg(cell, "1E3A8A")
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(label)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        # Données
        for row_idx, flow in enumerate(flows_q):
            row = table.add_row()
            if row_idx % 2 == 1:
                for cell in row.cells:
                    set_cell_bg(cell, "EFF6FF")
            for j, (_, (_, getter)) in enumerate(active_cols):
                cell = row.cells[j]
                cell.text = getter(flow)
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].font.size = Pt(8)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"flux_{ts}.docx"
    return Response(
        content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/{flow_id}")
def get_flow(flow_id: int, db: Session = Depends(get_db)):
    flow = db.query(FlowRequest).filter(FlowRequest.id == flow_id).first()
    if not flow:
        raise HTTPException(status_code=404, detail="Flux non trouvé")
    return {
        "id": flow.id,
        "created_at": flow.created_at.isoformat(),
        "src_ip": flow.src_ip,
        "dst_ip": flow.dst_ip,
        "port": flow.port,
        "protocol": flow.protocol,
        "application": flow.application,
        "justification": flow.justification,
        "status": flow.status,
        "analyst": flow.analyst,
        "rejection_reason": flow.rejection_reason,
        "validation": json.loads(flow.validation_result or "{}"),
        "path": json.loads(flow.path_result or "{}"),
        "scripts": json.loads(flow.scripts_result or "{}"),
    }


@router.patch("/{flow_id}/status")
def update_flow_status(flow_id: int, data: StatusUpdate, request: Request = None, db: Session = Depends(get_db)):
    flow = db.query(FlowRequest).filter(FlowRequest.id == flow_id).first()
    if not flow:
        raise HTTPException(status_code=404, detail="Flux non trouvé")
    allowed = {"pending", "validated", "deployed", "rejected"}
    if data.status not in allowed:
        raise HTTPException(status_code=400, detail=f"Statut invalide. Valeurs : {allowed}")
    old_status = flow.status
    flow.status = data.status
    if data.status == "rejected" and data.rejection_reason is not None:
        flow.rejection_reason = data.rejection_reason
    elif data.status != "rejected":
        flow.rejection_reason = None
    db.commit()

    # VALIDATE pour validated/deployed, sinon UPDATE ; FAILURE si rejet
    action = "VALIDATE" if data.status in ("validated", "deployed") else "UPDATE"
    audit.record_audit(
        db, action=action, object_type="FLOW",
        object_id=f"FLOW-{flow_id}",
        object_name=f"{flow.src_ip} → {flow.dst_ip}:{flow.port}",
        status="FAILURE" if data.status == "rejected" else "SUCCESS",
        application=flow.application or None,
        before={"status": old_status},
        after={"status": data.status, **({"rejection_reason": data.rejection_reason} if data.status == "rejected" else {})},
        ip_address=audit.client_ip(request),
    )
    return {"id": flow_id, "status": data.status}


@router.delete("/{flow_id}", status_code=204)
def delete_flow(flow_id: int, request: Request = None, db: Session = Depends(get_db)):
    flow = db.query(FlowRequest).filter(FlowRequest.id == flow_id).first()
    if not flow:
        raise HTTPException(status_code=404, detail="Flux non trouvé")
    snapshot = {"src_ip": flow.src_ip, "dst_ip": flow.dst_ip, "port": flow.port, "status": flow.status}
    label = f"{flow.src_ip} → {flow.dst_ip}:{flow.port}"
    app_name = flow.application or None
    db.delete(flow)
    db.commit()
    audit.record_audit(
        db, action="DELETE", object_type="FLOW",
        object_id=f"FLOW-{flow_id}", object_name=label, application=app_name,
        before=snapshot, ip_address=audit.client_ip(request),
    )


@router.get("/audit/summary")
def audit_summary(db: Session = Depends(get_db)):
    total = db.query(FlowRequest).count()
    validated = db.query(FlowRequest).filter(FlowRequest.status == "validated").count()
    deployed = db.query(FlowRequest).filter(FlowRequest.status == "deployed").count()
    rejected = db.query(FlowRequest).filter(FlowRequest.status == "rejected").count()
    pending = db.query(FlowRequest).filter(FlowRequest.status == "pending").count()

    # Top applications
    all_flows = db.query(FlowRequest).all()
    app_counts = {}
    for f in all_flows:
        app = f.application or "Non spécifiée"
        app_counts[app] = app_counts.get(app, 0) + 1
    top_apps = sorted(app_counts.items(), key=lambda x: -x[1])[:5]

    # Top analysts
    analyst_counts = {}
    for f in all_flows:
        a = f.analyst or "inconnu"
        analyst_counts[a] = analyst_counts.get(a, 0) + 1
    top_analysts = sorted(analyst_counts.items(), key=lambda x: -x[1])[:5]

    return {
        "total": total,
        "validated": validated,
        "deployed": deployed,
        "rejected": rejected,
        "pending": pending,
        "top_applications": [{"name": a, "count": c} for a, c in top_apps],
        "top_analysts": [{"name": a, "count": c} for a, c in top_analysts],
    }
